from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Election_Monitoring_Roles_and_Responsibilities_Party.docx"
LOGO = ROOT / "public" / "bsa-logo.png"

WINE = "741F3D"
DARK_WINE = "2B0816"
GOLD = "D9AA4B"
LIGHT_GOLD = "F5DC9A"
CREAM = "FFF8E8"
INK = "261018"
MUTED = "6B5560"
PALE = "F8EFF2"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ("Title", 28, DARK_WINE, 0, 8),
    ("Heading 1", 18, WINE, 15, 7),
    ("Heading 2", 13, DARK_WINE, 11, 5),
    ("Heading 3", 11, WINE, 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = name != "Title"
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Role Label" not in styles:
    role_style = styles.add_style("Role Label", WD_STYLE_TYPE.PARAGRAPH)
else:
    role_style = styles["Role Label"]
role_style.font.name = "Calibri"
role_style.font.size = Pt(9)
role_style.font.bold = True
role_style.font.color.rgb = RGBColor.from_string(GOLD)
role_style.paragraph_format.space_after = Pt(4)

bullet = styles["List Bullet"]
bullet.font.name = "Calibri"
bullet.font.size = Pt(10.5)
bullet.paragraph_format.left_indent = Inches(0.32)
bullet.paragraph_format.first_line_indent = Inches(-0.18)
bullet.paragraph_format.space_after = Pt(3)
bullet.paragraph_format.line_spacing = 1.15

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def bullet_item(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

def role_section(title, purpose, can_see, actions, limits=None):
    p = doc.add_paragraph(style="Role Label")
    p.add_run("SYSTEM ROLE")
    doc.add_heading(title, level=1)
    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(8)
    r = lead.add_run(purpose)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_WINE)
    doc.add_heading("What this role can see", level=2)
    for item in can_see: bullet_item(item)
    doc.add_heading("Main responsibilities and actions", level=2)
    for item in actions: bullet_item(item)
    if limits:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.7)
        cell = table.cell(0, 0)
        shade(cell, "F9E8ED"); set_cell_margins(cell, 120, 150, 120, 150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run("Access boundary: "); rr.bold = True; rr.font.color.rgb = RGBColor.from_string(WINE)
        p.add_run(limits)

# Running header/footer
header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("ELECTION MONITORING COMMAND CENTER  |  ROLE GUIDE")
hr.font.name = "Calibri"; hr.font.size = Pt(8); hr.font.bold = True; hr.font.color.rgb = RGBColor.from_string(WINE)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Operational reference  |  Oyo State Election Monitoring")
fr.font.name = "Calibri"; fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string(MUTED)

# Cover / opening block
if LOGO.exists():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(2.4))
p = doc.add_paragraph(); p.style = styles["Role Label"]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("OPERATIONS REFERENCE GUIDE")
p = doc.add_paragraph(); p.style = styles["Title"]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Roles and Responsibilities")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Election Monitoring Command Center")
r.font.size = Pt(14); r.bold = True; r.font.color.rgb = RGBColor.from_string(WINE)
p.paragraph_format.space_after = Pt(10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("This guide explains what each account type can see, what actions it performs, and where its access ends.")
r.font.size = Pt(10.5); r.font.color.rgb = RGBColor.from_string(MUTED)
p.paragraph_format.space_after = Pt(16)

doc.add_heading("Role overview", level=1)
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [1.28, 1.60, 2.38, 1.42]
headers = ["Role", "Operational scope", "Primary purpose", "Result reporting"]
for i, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
    cell.width = Inches(widths[i]); shade(cell, WINE); set_cell_margins(cell)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.bold = True; run.font.color.rgb = RGBColor.from_string(WHITE); run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
set_repeat_header(table.rows[0])
rows = [
    ("Super Admin", "Entire system", "Hidden technical and highest-level administration", "View all"),
    ("Admin", "Entire command", "Command-center monitoring, management and decisions", "View party total and all units"),
    ("Response Team", "Assigned response work", "Receive assignments and coordinate field response", "View assigned reports"),
    ("Supervisor", "Assigned LGA and ward", "Supervise agents and reports within one zone", "Submit for units in ward"),
    ("Agent", "Assigned polling unit", "Capture field evidence, location, incidents and results", "Submit assigned unit once"),
]
for idx, values in enumerate(rows):
    cells = table.add_row().cells
    for i, (cell, text) in enumerate(zip(cells, values)):
        cell.width = Inches(widths[i]); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if idx % 2: shade(cell, PALE)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text); run.font.size = Pt(8.7)
        if i == 0: run.bold = True; run.font.color.rgb = RGBColor.from_string(WINE)

doc.add_paragraph()
callout = doc.add_table(rows=1, cols=1); callout.alignment = WD_TABLE_ALIGNMENT.CENTER; callout.autofit = False
cell = callout.cell(0, 0); shade(cell, CREAM); set_cell_margins(cell, 130, 160, 130, 160)
p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
r = p.add_run("Core control principle: "); r.bold = True; r.font.color.rgb = RGBColor.from_string(WINE)
p.add_run("Admin and Super Admin have command-wide visibility. Supervisor is limited to an assigned LGA and ward. Agent is limited to an assigned polling unit and does not receive the operational map.")

doc.add_page_break()
role_section(
    "Super Admin",
    "The hidden highest-privilege account used for system ownership, recovery and exceptional administration.",
    ["Every SOS alert, incident, polling-unit result and uploaded evidence.", "All users, locations, camera shares, wards, polling units, map layers and command analytics.", "The same command-wide Results Center available to Admin."],
    ["Create and manage Admin, Response Team, Supervisor and Agent accounts.", "Maintain system-level settings, map data and protected administrative access.", "Recover access, correct exceptional configuration problems and supervise the full platform.", "Review the combined party vote total and the table of party results from every reporting polling unit."],
    "This account is intentionally hidden from ordinary user lists and should be used only for system-level work."
)

role_section(
    "Admin",
    "The command-center role responsible for the complete operational picture and election-result progress.",
    ["All SOS alerts and incidents from every zone.", "All party polling-unit results, vote counts and photographic evidence.", "All Agent and Supervisor GPS locations and active shared camera feeds.", "Command analytics, map information, personnel and the full Results Center."],
    ["Monitor incidents and SOS alerts across the election operation.", "Open the Results page to see the combined party total at the top and the polling-unit result table below.", "Review evidence images, LGA, ward, polling unit, upload time and party votes.", "Create and assign Response Team, Supervisor and Agent accounts with correct LGA, ward and polling-unit details.", "Coordinate response activity, update incident status and manage operational information."],
    "Admin visibility is command-wide. It is not restricted by ward, LGA or polling unit."
)

doc.add_page_break()
role_section(
    "Response Team",
    "The operational response role that receives and acts on assigned incidents and emergency information.",
    ["Incidents assigned or explicitly shared with the Response Team user.", "Relevant incident chat and operational information needed for the response.", "Nearby or operationally relevant SOS notifications when routed by the system."],
    ["Acknowledge and respond to assigned incidents.", "Share GPS/location and camera information when field visibility is needed.", "Coordinate through incident communication and report response progress.", "Use permitted field reporting functions for verified operational observations."],
    "Response Team members do not receive command-wide personnel management or full election-result administration."
)

role_section(
    "Supervisor",
    "The ward-level coordinator responsible for agents and reports inside one assigned operational zone.",
    ["The map for the Supervisor's assigned LGA and ward.", "Incidents, SOS alerts and results belonging to that assigned zone.", "Agents assigned to the same LGA and ward and their relevant live field positions."],
    ["Use the restricted sidebar to report a Polling Result or Incident.", "Submit results for polling units within the assigned ward, including text vote counts and mandatory signed-result photo evidence.", "Share GPS/location, share camera, share the zone map and send SOS.", "Monitor the zone's incident list and coordinate agents operating in the ward."],
    "A Supervisor cannot view other wards or command-wide analytics and management tools. Duplicate polling-unit results are rejected."
)

doc.add_page_break()
role_section(
    "Agent",
    "The polling-unit field role responsible for verified, timely evidence from one assigned polling unit.",
    ["A dedicated field dashboard rather than the operational map.", "The Agent's assigned LGA, ward and polling unit.", "The Agent's own reports and relevant operational feedback."],
    ["Report the party vote count for the assigned polling unit as one whole number.", "Attach a clear photograph of the signed polling-unit result; submission is blocked without picture evidence.", "Share live GPS/location with command.", "Share the phone camera and view the same live preview being transmitted.", "Switch between the front and back phone cameras while sharing.", "Send SOS with the current location and report urgent field incidents when permitted."],
    "An Agent does not see the operational map, command analytics, other zones or personnel-management tools. The Agent is locked to the assigned polling unit, and only one result is accepted for that unit."
)

doc.add_heading("Result reporting and verification workflow", level=1)
for step in [
    "Admin assigns each Supervisor to an LGA and ward, and each Agent to an LGA, ward and polling unit.",
    "Agent or Supervisor selects Polling Result, enters the party vote count and attaches the signed-result photograph.",
    "The server verifies the assignment, requires picture evidence and rejects a second result for the same LGA, ward and polling unit.",
    "Admin opens Results to monitor the number of reporting units, the combined party total and the detailed polling-unit table.",
    "If a correction is required, the field user contacts Admin rather than creating a contradictory duplicate submission.",
]:
    bullet_item(step)

doc.add_heading("Recommended account-control rules", level=1)
for item in [
    "Create accounts only after confirming the user's actual role, LGA, ward and polling unit.",
    "Do not share Admin or Super Admin credentials with field users.",
    "Use the minimum role required for the person's work.",
    "Confirm camera and location permissions on field phones before election-day deployment.",
    "Treat uploaded result photographs as election evidence and retain access only for authorized personnel.",
]:
    bullet_item(item)

doc.core_properties.title = "Election Monitoring Command Center - Roles and Responsibilities"
doc.core_properties.subject = "Operational role permissions and responsibilities"
doc.core_properties.author = "Election Monitoring Command Center"
doc.save(OUT)
print(OUT)
